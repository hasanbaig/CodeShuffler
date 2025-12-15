import copy
import random
import re

from docx import Document
from docx2python import docx2python
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def _create_border(border_type, val, sz=4, space=4, color="auto"):
    b = OxmlElement(f"w:{border_type}")
    b.set(qn("w:val"), val)
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), color)
    return b


def _add_code_block(doc, code_text: str):
    para = doc.add_paragraph()
    pPr = para._element.get_or_add_pPr()
    pShading = OxmlElement("w:shd")
    pShading.set(qn("w:val"), "clear")
    pShading.set(qn("w:color"), "auto")
    pShading.set(qn("w:fill"), "EAEAEA")
    pPr.insert(0, pShading)
    pBdr = OxmlElement("w:pBdr")

    pBdr.append(_create_border("top", "single", color="000000", sz=6))
    pBdr.append(_create_border("left", "single", color="000000", sz=6))
    pBdr.append(_create_border("bottom", "single", color="000000", sz=6))
    pBdr.append(_create_border("right", "single", color="000000", sz=6))

    pPr.append(pBdr)
    hidden_start = para.add_run("<code>")
    hidden_start.font.hidden = True

    lines = code_text.split("\\n")

    for idx, line in enumerate(lines):
        stripped = line.lstrip(" ")
        leading_spaces = len(line) - len(stripped)
        visible_line = ("\u00A0" * leading_spaces) + stripped

        run = para.add_run(visible_line)
        run.font.name = "Courier New"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        if idx < len(lines) - 1:
            para.add_run("\n")
    hidden_end = para.add_run("</code>")
    hidden_end.font.hidden = True
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    return para


def create_exam_docx(template_path, exam_dict, output_path, answer_key=False):
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Error loading template file {template_path}: {e}")
        return
    _append_questions(doc, exam_dict, answer_key=answer_key)
    doc.save(output_path)
    print(f"Exam rebuilt and saved to {output_path}")


def _append_questions(doc, exam_dict, answer_key):
    for qnum, content in exam_dict.items():
        question = content["question"]
        code_lines = content.get("code", [])
        choices = content["choices"]
        correct_index = content.get("correct_index")

        q_para = doc.add_paragraph()
        run = q_para.add_run(f"{qnum}) {question}")
        run.font.size = Pt(11)
        q_para.paragraph_format.space_after = Pt(6)

        if code_lines:
            code_text = "\\n".join(code_lines)
            _add_code_block(doc, code_text)

        for i, choice in enumerate(choices):
            letter = chr(97 + i)
            c_para = doc.add_paragraph()
            c_para.paragraph_format.left_indent = Pt(18)
            c_para.paragraph_format.space_after = Pt(3)

            run = c_para.add_run(f"({letter}) {choice}")
            if answer_key and correct_index == i:
                run.bold = True

        doc.add_paragraph("")


def parse_exam(doc_path: str):
    with docx2python(doc_path) as docx_content:
        text = docx_content.text

    lines = []
    for raw in text.splitlines():
        stripped = raw.strip()
        if stripped:
            lines.append(stripped)

    exam = {}
    current_qnum = None
    choices: list[str] = []
    question_text: list[str] = []
    code_buffer: list[str] = []
    correct_index = None
    in_code = False
    question_pattern = re.compile(r"^(\d+)[).]")
    choice_pattern = re.compile(r"^\(?[a-dA-D]\)?[).]")

    for line in lines:
        # case 1: code starts and ends on the same line
        if "<code>" in line and "</code>" in line:
            before, rest = line.split("<code>", 1)
            code_part, after = rest.split("</code>", 1)
            if before.strip() and current_qnum is not None and not question_pattern.match(line):
                question_text.append(before.strip())
            code_part = code_part.strip()
            if code_part:
                code_buffer.append(code_part)
            in_code = False
            if after.strip() and current_qnum is not None:
                question_text.append(after.strip())
            continue

        # case 2: code starts on this line (possibly mid-line)
        if "<code>" in line:
            before, after = line.split("<code>", 1)
            if before.strip() and current_qnum is not None:
                question_text.append(before.strip())
            code_part = after.strip()
            if code_part:
                code_buffer.append(code_part)
            in_code = True
            continue

        # case 3: code ends on this line (possibly mid-line)
        if "</code>" in line:
            before, after = line.split("</code>", 1)
            code_part = before.strip()
            if code_part:
                code_buffer.append(code_part)
            in_code = False
            if after.strip() and current_qnum is not None:
                question_text.append(after.strip())
            continue
        if in_code:
            code_buffer.append(line)
            continue
        if question_pattern.match(line):
            if current_qnum is not None:
                exam[current_qnum] = {
                    "question": " ".join(question_text).strip(),
                    "code": code_buffer,
                    "choices": choices,
                    "correct_index": correct_index,
                }
            m = question_pattern.match(line)
            if m is not None:
                current_qnum = int(m.group(1))
                question_text = [re.sub(question_pattern, "", line).strip()]
                choices = []
                code_buffer = []
                correct_index = None
            continue

        if choice_pattern.match(line) and current_qnum is not None:
            choice_text = re.sub(choice_pattern, "", line).strip()
            if choice_text.endswith("*"):
                correct_index = len(choices)
                choice_text = choice_text[:-1].rstrip()
            choices.append(choice_text)
            continue

        if current_qnum is not None:
            question_text.append(line)

    if current_qnum is not None:
        exam[current_qnum] = {
            "question": " ".join(question_text).strip(),
            "code": code_buffer,
            "choices": choices,
            "correct_index": correct_index,
        }
    return exam


def shuffle_questions(exam_dict):
    exam_copy = copy.deepcopy(exam_dict)
    question_items = list(exam_copy.items())
    random.shuffle(question_items)

    shuffled_exam = {}
    for new_qnum, (_, content) in enumerate(question_items, start=1):
        shuffled_exam[new_qnum] = content

    return shuffled_exam


def shuffle_answers(exam_dict):
    exam_copy = copy.deepcopy(exam_dict)

    for _, content in exam_copy.items():
        choices = content.get("choices")
        indexed = list(enumerate(choices))

        random.shuffle(indexed)
        new_choices = [text for (_, text) in indexed]
        old_correct = content.get("correct_index")
        new_correct = None

        if old_correct is not None:
            for new_i, (old_i, _) in enumerate(indexed):
                if old_i == old_correct:
                    new_correct = new_i
                    break

        content["choices"] = new_choices
        content["correct_index"] = new_correct

    return exam_copy


def print_exam_dict(exam_dict):
    print("\n" + "=" * 80)
    print(f"EXAM CONTENTS - {len(exam_dict)} Questions")
    print("=" * 80 + "\n")

    for question_num in sorted(exam_dict.keys()):
        question_data = exam_dict[question_num]
        print(f"Question {question_num}:")
        print(f"  {question_data['question']}")
        print()
        labels = ["a", "b", "c", "d", "e", "f"]
        print("  Choices:")
        for idx, choice in enumerate(question_data["choices"]):
            if idx < len(labels):
                print(f"    ({labels[idx]}) {choice}")
            else:
                print(f"    - {choice}")
        print("\n" + "-" * 80 + "\n")
